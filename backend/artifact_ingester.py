"""
Artifact ingestion — reads PDF / DOCX / plain text, sends to Claude,
returns a structured scheme JSON ready to drop into schemes/.

The system decides. AI extracts — then a human reviews before saving.
"""

import io
import json
import os
import re
import anthropic

EXTRACTION_SYSTEM_PROMPT = """You are a NSW Government procurement policy analyst.
Your job is to read procurement policy documents and extract their rules into a structured JSON format.

You must output ONLY valid JSON. No explanation, no markdown, no prose outside the JSON.

The JSON must conform exactly to this schema:

{
  "scheme_id": "UPPERCASE_SNAKE_CASE identifier (e.g. ICT_WOG, CONSTRUCTION, LABOUR_HIRE)",
  "name": "Full human-readable scheme name",
  "version": "YYYY-MM or version string from the document",
  "active": true,
  "description": "One paragraph describing what this scheme covers and when it applies",
  "source": "Full citation — Act name, Board Direction number, Policy name and date",
  "triggers": {
    "category": ["ict_saas","ict_hardware","professional_services","consulting","goods","construction","labour_hire","other"],
    "value": ["micro","low","medium","high","major"],
    "purpose": ["new","renewal","emergency","pilot","replacement"],
    "market": ["sole","limited","some","broad","unknown"],
    "impact": ["low","medium","high","critical"],
    "overlays": ["ai","privacy","critical_ict","construction","overseas"],
    "timing": ["urgent","compressed","normal","extended","unknown"],
    "org": ["operational","corporate","executive","central"],
    "interaction": ["minimal","quotes","tender","collaborative"]
  },
  "pre_checks": [
    {
      "id": "SCHEMEID-PC-001",
      "title": "Short action title",
      "body": "Plain English — what to check before proceeding",
      "link": null,  // URL if available, otherwise null
      "citation": "Exact section, clause, or page reference in the source document — e.g. 'Section 4.2, p.15' or 'Clause 3(1)(a)'"
    }
  ],
  "obligations": [
    {
      "id": "SCHEMEID-OBL-001",
      "title": "Short obligation title",
      "body": "Plain English description of what must be done — no legalese",
      "policy": "Exact policy/Act reference with section number",
      "citation": "Exact section, clause, or page reference in the source document — e.g. 'Section 4.2, p.15' or 'Clause 3(1)(a)'"
    }
  ],
  "pathway_overrides": {
    "if_arrangement_exists": "What to do if a scheme arrangement already covers this (or null)",
    "if_no_arrangement": "What to do if no arrangement exists (or null)"
  },
  "step_injections": [
    {
      "after_step": 0,
      "steps": [
        {
          "text": "Step text to inject at the start of the process",
          "citation": "Exact section, clause, or page reference in the source document"
        }
      ]
    }
  ],
  "approval_additions": [
    {
      "role": "Name of additional required approver or reviewer",
      "note": "Why they are required"
    }
  ]
}

Rules for extraction:
0. Output ONLY a single valid JSON object. No comments, no trailing commas, no markdown.
1. Extract ALL substantive obligations — thresholds, mandatory actions, prohibited actions, review requirements
2. Convert dollar thresholds into the value tier system (micro/low/medium/high/major)
3. Write obligations in plain English — translate legal language into what a public servant needs to do
4. Identify the correct triggers — be specific about which procurement categories and value ranges apply
5. If a threshold applies above a certain value, include all tiers at or above that value
6. If the document covers multiple distinct obligation areas, include each as a separate obligation
7. If you cannot determine a value for a field, use null for strings or [] for arrays
8. The scheme_id must be unique and descriptive — use the policy acronym if available
9. For every obligation, pre-check, and step injection step, populate the citation field with the precise section number, clause, heading, or page reference from the source document where that rule originates — this is used to trace determinations back to source material. If no specific location can be identified, use the document title and section heading."""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_text_from_file(filename: str, file_bytes: bytes) -> str:
    """Dispatch to appropriate extractor based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    elif ext in ("txt", "md"):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        # Try UTF-8 decode as fallback
        return file_bytes.decode("utf-8", errors="replace")


def ingest_artifact(
    pasted_text: str = "",
    filename: str = "",
    file_bytes: bytes = b"",
) -> dict:
    """
    Main ingestion function.
    Accepts pasted text and/or a file. Sends combined content to Claude.
    Returns a dict with:
      - scheme: the extracted scheme JSON (dict)
      - source_text_preview: first 500 chars of what was extracted
      - token_count: approximate
      - warnings: list of any issues detected
    """
    if not os.environ.get("ANTHROPIC_API_KEY"):
        raise ValueError("ANTHROPIC_API_KEY not configured")

    # Build source text
    source_parts = []
    file_text = ""

    if file_bytes and filename:
        try:
            file_text = extract_text_from_file(filename, file_bytes)
            if file_text.strip():
                source_parts.append(f"=== DOCUMENT: {filename} ===\n{file_text}")
        except Exception as e:
            raise ValueError(f"Could not read file '{filename}': {e}")

    if pasted_text and pasted_text.strip():
        source_parts.append(f"=== PASTED TEXT ===\n{pasted_text.strip()}")

    if not source_parts:
        raise ValueError("No content provided — supply pasted text and/or a file.")

    combined = "\n\n".join(source_parts)

    # Truncate to avoid context limits (~80k chars ≈ ~20k tokens, well within Claude's limit)
    if len(combined) > 80_000:
        combined = combined[:80_000] + "\n\n[... document truncated for processing ...]"

    client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

    user_prompt = f"""Extract the procurement rules from the following policy document(s) into the required JSON schema.

{combined}

Return ONLY the JSON object. No markdown code fences, no explanation."""

    response = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=4096,
        system=[{
            "type": "text",
            "text": EXTRACTION_SYSTEM_PROMPT,
            "cache_control": {"type": "ephemeral"},
        }],
        messages=[{"role": "user", "content": user_prompt}],
    )

    raw = response.content[0].text.strip()

    # Strip markdown fences if present
    if raw.startswith("```"):
        lines = raw.split("\n")
        raw = "\n".join(lines[1:])
        if raw.endswith("```"):
            raw = raw.rsplit("```", 1)[0]
    raw = raw.strip()

    warnings = []

    from json_repair import repair_json
    try:
        repaired = repair_json(raw, return_objects=True)
        if not isinstance(repaired, dict):
            raise ValueError(f"Extracted content is not a JSON object (got {type(repaired).__name__})")
        scheme = repaired
    except Exception as e:
        raise ValueError(f"Could not parse extracted JSON: {e}")

    # Validate required fields
    required = ["scheme_id", "name", "triggers", "obligations"]
    missing = [f for f in required if f not in scheme]
    if missing:
        warnings.append(f"Missing recommended fields: {', '.join(missing)}")

    # Ensure active is set
    scheme.setdefault("active", True)

    # Ensure arrays are arrays
    for arr_field in ["pre_checks", "obligations", "step_injections", "approval_additions"]:
        if arr_field not in scheme or not isinstance(scheme[arr_field], list):
            scheme[arr_field] = []

    if not scheme.get("triggers"):
        warnings.append("No triggers were identified — review and set manually before activating.")

    if len(scheme.get("obligations", [])) == 0:
        warnings.append("No obligations were extracted — document may not contain actionable procurement rules.")

    source_preview = combined[:400] + ("..." if len(combined) > 400 else "")

    return {
        "scheme": scheme,
        "source_text_preview": source_preview,
        "char_count": len(combined),
        "warnings": warnings,
    }
