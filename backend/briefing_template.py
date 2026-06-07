from __future__ import annotations

from datetime import date
from io import BytesIO
import os
from pathlib import Path
import re
import zipfile

from docx import Document

from models import ApprovalResult, BriefingStructure, PathwayResult, ProcurementProfile


TEMPLATE_NAME = "DCS_Briefing Note 1.dotx"
DEFAULT_TEMPLATE_PATHS = [
    Path(__file__).parent / "templates" / TEMPLATE_NAME,
    Path(r"C:\Users\garat\Downloads\DCS_Briefing Note 1.dotx"),
]
TEMPLATE_HEADINGS = {
    "Recommendations",
    "Purpose",
    "Background",
    "Consultation (delete if not applicable)",
    "Attachments (delete if not applicable)",
    "Approval",
}
TEMPLATE_BLOCK_STOP_LINES = {
    "Signature:",
    "Contact Officer:",
}


def render_briefing_note(
    briefing_structure: BriefingStructure,
    prose_sections: dict[str, str],
    profile: ProcurementProfile,
    pathway: PathwayResult,
    approvals: ApprovalResult,
) -> tuple[str, bytes]:
    _ = briefing_structure
    template_path = _resolve_template_path()
    doc = Document(_dotx_to_docx_stream(template_path.read_bytes()))

    title = _build_title(profile, pathway)
    issue = _build_issue_summary(prose_sections)
    recommendations = _build_recommendations(prose_sections, pathway)
    purpose = prose_sections.get("purpose", "").strip()
    background = _build_background(prose_sections)
    consultation_note = _build_consultation_note(approvals)
    attachments_text = _build_attachments_text()
    contact_officer = _build_contact_officer(profile)
    approver_rows = _build_approver_rows(approvals)

    doc.core_properties.title = title
    _set_first_issue_table(doc, issue)
    _set_recommendations_block(doc, recommendations)
    _insert_content_before_heading(doc, "Background", purpose)
    _insert_content_before_heading(doc, "Consultation (delete if not applicable)", background)
    _insert_content_before_heading(doc, "Attachments (delete if not applicable)", consultation_note)
    _insert_content_before_heading(doc, "Approval", attachments_text)
    _set_contact_officer(doc, contact_officer)
    _set_approval_table(doc, approver_rows)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return f"{_slugify(title) or 'briefing-note'}.docx", out.getvalue()


def get_template_prompt_guidance() -> dict[str, dict]:
    """
    Extract concise authoring guidance from the live briefing note template so it can
    be injected into the LLM prompt. This keeps drafting instructions aligned with
    the actual .dotx used for export.
    """
    template_path = _resolve_template_path()
    doc = Document(_dotx_to_docx_stream(template_path.read_bytes()))
    paragraph_blocks = _extract_template_heading_blocks(doc)

    return {
        "purpose": {
            "template_anchor": "Background",
            "placement": "Rendered immediately before the 'Background' heading in the template.",
        },
        "background": {
            "template_anchor": "Consultation (delete if not applicable)",
            "placement": "Rendered into the main body before the consultation section.",
        },
        "what_is_being_procured": {
            "template_anchor": "Consultation (delete if not applicable)",
            "placement": "Rendered into the main body before the consultation section.",
        },
        "recommended_pathway": {
            "template_anchor": "Consultation (delete if not applicable)",
            "placement": "Rendered into the main body before the consultation section.",
        },
        "value_for_money": {
            "template_anchor": "Consultation (delete if not applicable)",
            "placement": "Rendered into the main body before the consultation section.",
        },
        "market_context": {
            "template_anchor": "Consultation (delete if not applicable)",
            "placement": "Rendered into the main body before the consultation section.",
        },
        "risk_and_impact": {
            "template_anchor": "Consultation (delete if not applicable)",
            "placement": "Rendered into the main body before the consultation section.",
        },
        "obligations": {
            "template_anchor": "Consultation (delete if not applicable)",
            "placement": "Rendered into the main body before the consultation section.",
        },
        "approvals": {
            "template_anchor": "Consultation (delete if not applicable)",
            "placement": "Rendered into the main body before the consultation section.",
            "approval_table_headers": _get_table_headers(doc, 2),
        },
        "dcs_compliance": {
            "template_anchor": "Consultation (delete if not applicable)",
            "placement": "Rendered into the main body before the consultation section.",
        },
        "justification": {
            "template_anchor": "Consultation (delete if not applicable)",
            "placement": "Rendered into the main body before the consultation section.",
        },
        "recommendation": {
            "template_anchor": "Recommendations",
            "instructions": paragraph_blocks.get("Recommendations", []),
        },
        "issue_summary": {
            "template_table_headers": _get_table_headers(doc, 0),
        },
        "consultation_table": {
            "template_table_headers": _get_table_headers(doc, 1),
        },
    }


def _resolve_template_path() -> Path:
    env_path = os.environ.get("BRIEFING_TEMPLATE_PATH")
    if env_path:
        candidate = Path(env_path)
        if candidate.exists():
            return candidate

    for candidate in DEFAULT_TEMPLATE_PATHS:
        if candidate.exists():
            return candidate

    raise FileNotFoundError(
        "Briefing note template not found. Set BRIEFING_TEMPLATE_PATH or place the template at "
        f"{DEFAULT_TEMPLATE_PATHS[0]}"
    )


def _extract_template_heading_blocks(doc: Document) -> dict[str, list[str]]:
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    blocks: dict[str, list[str]] = {}

    for idx, text in enumerate(texts):
        if text not in TEMPLATE_HEADINGS:
            continue

        block: list[str] = []
        cursor = idx + 1
        while (
            cursor < len(texts)
            and texts[cursor] not in TEMPLATE_HEADINGS
            and texts[cursor] not in TEMPLATE_BLOCK_STOP_LINES
        ):
            block.append(texts[cursor])
            cursor += 1
        blocks[text] = block

    return blocks


def _get_table_headers(doc: Document, index: int) -> list[str]:
    if index >= len(doc.tables):
        return []
    header_row = doc.tables[index].rows[0]
    return [cell.text.strip() for cell in header_row.cells if cell.text.strip()]


def _dotx_to_docx_stream(template_bytes: bytes) -> BytesIO:
    src = BytesIO(template_bytes)
    dst = BytesIO()
    with zipfile.ZipFile(src) as zin, zipfile.ZipFile(dst, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(
                    b"application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                    b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                )
            zout.writestr(item, data)
    dst.seek(0)
    return dst


def _build_title(profile: ProcurementProfile, pathway: PathwayResult) -> str:
    purpose = (profile.purpose or "procurement").replace("_", " ")
    category = (profile.category or "goods and services").replace("_", " ")
    purpose = purpose[:1].upper() + purpose[1:]
    category = category.replace("ict", "ICT")
    return f"{purpose} procurement for {category} via {pathway.label}"


def _build_issue_summary(prose_sections: dict[str, str]) -> str:
    seed = " ".join(
        filter(
            None,
            [
                prose_sections.get("purpose", ""),
                prose_sections.get("recommended_pathway", ""),
                prose_sections.get("recommendation", ""),
            ],
        )
    )
    sentences = _split_sentences(seed)
    if not sentences:
        return "Approval is sought to proceed with the recommended procurement pathway."
    return " ".join(sentences[:2]).strip()


def _build_recommendations(prose_sections: dict[str, str], pathway: PathwayResult) -> str:
    candidate = prose_sections.get("recommendation", "").strip()
    sentences = _split_sentences(candidate)
    lines: list[str] = []

    if sentences:
        first = sentences[0]
        if not re.match(r"^(Approve|Note|Oppose|Sign)\b", first, flags=re.IGNORECASE):
            first = (
                f"Approve {first[:1].lower() + first[1:]}"
                if first
                else f"Approve the recommended procurement pathway of {pathway.label}."
            )
        lines.append(first)

        for extra in sentences[1:3]:
            if not re.match(r"^(Approve|Note|Oppose|Sign)\b", extra, flags=re.IGNORECASE):
                extra = f"Note {extra[:1].lower() + extra[1:]}"
            lines.append(extra)

    if not lines:
        lines = [
            f"Approve the recommended procurement pathway of {pathway.label}.",
            "Note the policy obligations, approvals, and compliance requirements set out in this briefing note.",
        ]

    return "\n".join(lines)


def _build_background(prose_sections: dict[str, str]) -> str:
    ordered_ids = [
        "background",
        "what_is_being_procured",
        "recommended_pathway",
        "value_for_money",
        "market_context",
        "risk_and_impact",
        "obligations",
        "approvals",
        "dcs_compliance",
        "justification",
    ]
    chunks = []
    for section_id in ordered_ids:
        text = prose_sections.get(section_id, "").strip()
        if text:
            heading = section_id.replace("_", " ").title()
            chunks.append(f"{heading}\n{text}")
    return "\n\n".join(chunks).strip()


def _build_consultation_note(approvals: ApprovalResult) -> str:
    if approvals.reviews_required:
        reviews = ", ".join(approvals.reviews_required)
        return f"Required consultation and review points identified through the assessment: {reviews}."
    return "No mandatory consultation points were identified by the rules assessment. Confirm business-specific consultation requirements before submission."


def _build_attachments_text() -> str:
    return "Tab A - Procurement assessment output and supporting decision record."


def _build_contact_officer(profile: ProcurementProfile) -> str:
    agency = profile.agency_name or "Business owner"
    return f"Contact Officer: {agency} representative - complete name, position and contact number before submission."


def _build_approver_rows(approvals: ApprovalResult) -> list[str]:
    rows = [approvals.approver_role]
    rows.extend(add.role for add in approvals.additions)
    rows.extend(review for review in approvals.reviews_required if review not in rows)
    return rows[:2] if rows else ["Approver to be confirmed"]


def _find_paragraph(doc: Document, anchor_text: str):
    paragraphs = doc.paragraphs
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == anchor_text:
            return paragraph
    raise ValueError(f"Could not find heading '{anchor_text}' in briefing template")


def _insert_content_before_heading(doc: Document, heading_text: str, content: str) -> None:
    anchor = _find_paragraph(doc, heading_text)
    lines = [line.strip() for line in content.split("\n") if line.strip()]
    if not lines:
        return
    for line in reversed(lines):
        anchor.insert_paragraph_before(line)


def _set_first_issue_table(doc: Document, issue: str) -> None:
    if not doc.tables:
        raise ValueError("Briefing template does not contain the issue table")

    table = doc.tables[0]
    row = table.rows[1]
    row.cells[0].text = f"Briefing\n\n{issue}"
    _set_signature_date(doc)


def _set_signature_date(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Signature:"):
            paragraph.text = f"Signature:    {_format_date(date.today())}"
            return


def _set_recommendations_block(doc: Document, recommendations: str) -> None:
    paragraphs = doc.paragraphs
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() == "Recommendations":
            if idx + 1 < len(paragraphs):
                paragraphs[idx + 1].text = recommendations
            if idx + 2 < len(paragraphs):
                paragraphs[idx + 2].text = ""
            return
    raise ValueError("Could not find Recommendations heading in briefing template")


def _set_contact_officer(doc: Document, contact_officer: str) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Contact Officer:"):
            paragraph.text = contact_officer
            return


def _set_approval_table(doc: Document, approver_rows: list[str]) -> None:
    if len(doc.tables) < 3:
        raise ValueError("Briefing template does not contain the approval table")

    table = doc.tables[2]
    for idx, approver in enumerate(approver_rows, start=1):
        if idx < len(table.rows):
            table.cell(idx, 0).text = approver
            table.cell(idx, 1).text = "[Enter a Date]"


def _split_sentences(text: str) -> list[str]:
    text = re.sub(r"\s+", " ", text or "").strip()
    if not text:
        return []
    return [sentence.strip() for sentence in re.split(r"(?<=[.!?])\s+", text) if sentence.strip()]


def _slugify(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")


def _format_date(value: date) -> str:
    return f"{value.day} {value.strftime('%B %Y')}"
